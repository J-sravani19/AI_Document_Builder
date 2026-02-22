from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
import re
from datetime import datetime

client = genai.Client()


def generate_cover_letter(name, phone, linkedin,
                          skills, role, company,
                          education, profile_type, template):

    today_date = datetime.today().strftime("%B %d, %Y")

    prompt = f"""
Write an ATS-friendly professional cover letter (300-400 words).

{name}
Phone: {phone}
LinkedIn: {linkedin}
Date: {today_date}

Hiring Manager
{company}

Subject: Application for {role}

Include:
- Strong introduction
- Skills aligned with job
- Education details: {education}
- Professional closing

Profile Type: {profile_type}
Skills: {skills}

Keep formatting simple.
Single column.
"""

    try:
        response = client.models.generate_content(
            model="models/gemini-2.5-flash",
            contents=prompt,
        )
        text = clean_text(response.text)

    except Exception:
        # Fallback ATS cover letter
        text = generate_ats_cover_letter(
            name, phone, linkedin, role,
            company, skills, education, today_date
        )

    safe_name = name.replace(" ", "_")
    docx_file = f"{safe_name}_CoverLetter.docx"
    pdf_file = f"{safe_name}_CoverLetter.pdf"

    save_docx(text, docx_file, template)
    save_pdf(text, pdf_file)

    return docx_file, pdf_file


def generate_ats_cover_letter(name, phone, linkedin,
                              role, company,
                              skills, education, date):

    return f"""
{name}
Phone: {phone}
LinkedIn: {linkedin}
Date: {date}

Hiring Manager
{company}

Subject: Application for {role}

Dear Hiring Manager,

I am writing to express my interest in the {role} position at {company}. With strong technical skills in {skills}, I am confident in my ability to contribute effectively to your team.

My educational background includes {education}, which has provided me with a solid foundation in analytical thinking, problem-solving, and technical expertise.

I am eager to apply my knowledge and skills in a dynamic environment and contribute to organizational success.

Thank you for considering my application.

Sincerely,
{name}
"""


def clean_text(text):
    text = re.sub(r'#+', '', text)
    text = re.sub(r'\*\*', '', text)
    return text.strip()


def save_docx(text, filename, template):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    first = True

    for line in text.split("\n"):
        clean = line.strip()
        if not clean:
            continue

        if first:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(16)
            if template == "Modern":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first = False
            continue

        doc.add_paragraph(clean)

    doc.save(filename)


def save_pdf(text, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    normal = ParagraphStyle(
        name="NormalStyle",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=6
    )

    for line in text.split("\n"):
        clean = line.strip()
        if clean:
            elements.append(Paragraph(clean, normal))
            elements.append(Spacer(1, 0.15 * inch))

    doc.build(elements)