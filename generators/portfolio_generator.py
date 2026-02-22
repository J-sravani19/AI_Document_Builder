from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
import re

client = genai.Client()


def generate_portfolio(name, phone, linkedin, github,
                       skills, projects,
                       certifications, achievements,
                       profile_type, template):

    prompt = f"""
Create a professional ATS-friendly portfolio document.

Name: {name}
Phone: {phone}
LinkedIn: {linkedin}
GitHub: {github}
Skills: {skills}
Projects: {projects}
Certifications: {certifications}
Achievements: {achievements}
Profile Type: {profile_type}

Structure:
ABOUT ME
SKILLS
PROJECTS
EDUCATION
CERTIFICATIONS
ACHIEVEMENTS

Use uppercase headings.
Use bullet points.
Single column layout.
Professional tone.
"""

    try:
        response = client.models.generate_content(
            model="models/gemini-2.5-flash",
            contents=prompt,
        )
        text = clean_text(response.text)

    except Exception:
        # Fallback ATS portfolio
        text = generate_ats_portfolio(
            name, phone, linkedin, github,
            skills, projects,
            certifications, achievements
        )

    safe_name = name.replace(" ", "_")
    docx_file = f"{safe_name}_Portfolio.docx"
    pdf_file = f"{safe_name}_Portfolio.pdf"

    save_docx(text, docx_file, template)
    save_pdf(text, pdf_file)

    return docx_file, pdf_file


def generate_ats_portfolio(name, phone, linkedin, github,
                           skills, projects,
                           certifications, achievements):

    return f"""
{name}

CONTACT
Phone: {phone}
LinkedIn: {linkedin}
GitHub: {github}

ABOUT ME
Motivated professional with strong technical and analytical skills.

SKILLS
• {skills}

PROJECTS
• {projects}

CERTIFICATIONS
• {certifications}

ACHIEVEMENTS
• {achievements}

EDUCATION
Bachelor's Degree in Relevant Field
University Name
Year of Graduation
"""


def clean_text(text):
    text = re.sub(r'#+', '', text)
    text = re.sub(r'\*\*', '', text)
    text = re.sub(r'^\s*[-*]\s+', '• ', text, flags=re.MULTILINE)
    return text.strip()


def save_docx(text, filename, template):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    first = True

    for line in text.split("\n"):
        clean = line.strip()
        if not clean:
            continue

        if first:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(18)
            if template == "Modern":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first = False
            continue

        if clean.isupper():
            h = doc.add_paragraph()
            run = h.add_run(clean)
            run.bold = True
            run.font.size = Pt(13)
            continue

        if clean.startswith("•"):
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(clean.replace("•", "").strip())
            continue

        doc.add_paragraph(clean)

    doc.save(filename)


def save_pdf(text, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    normal = ParagraphStyle(
        name="NormalStyle",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=6
    )

    for line in text.split("\n"):
        clean = line.strip()
        if clean:
            elements.append(Paragraph(clean, normal))
            elements.append(Spacer(1, 0.15 * inch))


    doc.build(elements)


