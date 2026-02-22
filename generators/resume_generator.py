from google import genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
import re

client = genai.Client()


def generate_resume(name, phone, linkedin, github,
                    skills, projects, role, job_description,
                    profile_type, certifications, template):

    prompt = f"""
Create a highly ATS-optimized resume.

Name: {name}
Phone: {phone}
LinkedIn: {linkedin}
GitHub: {github}
Target Role: {role}
Skills: {skills}
Projects: {projects}
Certifications: {certifications}
Job Description: {job_description}
Profile Type: {profile_type}

Use uppercase headings.
Use bullet points.
Single column format.
Include EDUCATION section.
"""

    try:
        response = client.models.generate_content(
            model="models/gemini-2.5-flash",
            contents=prompt,
        )
        text = clean_text(response.text)

    except Exception:
        # Fallback ATS Template (No AI)
        text = generate_ats_resume(
            name, phone, linkedin, github,
            skills, projects, certifications, role
        )

    safe_name = name.replace(" ", "_")
    docx_file = f"{safe_name}_Resume.docx"
    pdf_file = f"{safe_name}_Resume.pdf"

    save_docx(text, docx_file, template)
    save_pdf(text, pdf_file)

    return docx_file, pdf_file


def generate_ats_resume(name, phone, linkedin, github,
                        skills, projects, certifications, role):

    return f"""
{name}

CONTACT
Phone: {phone}
LinkedIn: {linkedin}
GitHub: {github}

PROFESSIONAL SUMMARY
Results-driven candidate seeking {role} position with strong technical and analytical skills.

SKILLS
• {skills}

PROJECTS
• {projects}

CERTIFICATIONS
• {certifications}

EDUCATION
Bachelor's Degree in Relevant Field
University Name
Year of Graduation
"""


def clean_text(text):
    text = re.sub(r'#+', '', text)
    text = re.sub(r'\*\*', '', text)
    text = re.sub(r'^\s*[-*]\s+', '• ', text, flags=re.MULTILINE)
    return text.strip()


def save_docx(text, filename, template):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    first = True

    for line in text.split("\n"):
        clean = line.strip()
        if not clean:
            continue

        if first:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(18)
            if template == "Modern":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first = False
            continue

        if clean.isupper():
            h = doc.add_paragraph()
            run = h.add_run(clean)
            run.bold = True
            run.font.size = Pt(13)
            continue

        if clean.startswith("•"):
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(clean.replace("•", "").strip())
            continue

        doc.add_paragraph(clean)

    doc.save(filename)


def save_pdf(text, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    normal = ParagraphStyle(
        name="NormalStyle",
        parent=styles["Normal"],
        fontSize=11,
        spaceAfter=6
    )

    for line in text.split("\n"):
        clean = line.strip()
        if clean:
            elements.append(Paragraph(clean, normal))
            elements.append(Spacer(1, 0.15 * inch))

    doc.build(elements)